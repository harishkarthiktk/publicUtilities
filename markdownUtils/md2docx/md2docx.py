#!/usr/bin/env python
"""
Markdown to DOCX Converter

Converts markdown files to Microsoft Word documents with configurable styling.
Supports headings, inline formatting (bold/italic/strikethrough/code), nested lists,
tables, code blocks, blockquotes, horizontal rules, definition lists, links,
Obsidian ==highlights==, and YAML frontmatter as DOCX metadata.

Usage:
    python md2docx.py input.md output.docx [--config config.yaml]
"""

import re
import sys
import argparse
import os
import logging
from html.parser import HTMLParser
import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
import mdformat

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# HTML Parser - extracts structured elements from markdown-generated HTML
# ---------------------------------------------------------------------------

class MarkdownHTMLParser(HTMLParser):
    """Parses HTML from markdown and extracts structured elements with proper
    per-span inline formatting."""

    def __init__(self):
        super().__init__()
        self.elements = []
        self.list_stack = []  # stack of "ul"/"ol"
        self.format_stack = []  # stack of active format tags: "bold", "italic", "code", "strikethrough", "mark"
        self.current_spans = []  # list of {"text": ..., "bold": bool, ...}
        self.current_type = None  # "h1"-"h6", "p", "li", etc.
        self.current_attrs = {}  # extra attrs like list_type, level, href

        # Table state
        self.in_table = False
        self.table_rows = []
        self.current_row = []
        self.current_cell_spans = []
        self.in_cell = False
        self.cell_tag = None  # "th" or "td"

        # Blockquote state
        self.in_blockquote = False

        # Pre/code block state
        self.in_pre = False
        self.pre_text = []

        # Link state
        self.current_href = None

        # Definition list state
        self.in_dl = False

        # Track which list levels need a counter restart on next li
        self._restart_next_li = set()

    def _current_format_state(self):
        return {
            "bold": "bold" in self.format_stack,
            "italic": "italic" in self.format_stack,
            "code": "code" in self.format_stack,
            "strikethrough": "strikethrough" in self.format_stack,
            "highlight": "mark" in self.format_stack,
        }

    def _flush_element(self):
        if self.current_type is None:
            return
        spans = [s for s in self.current_spans if s["text"]]
        if spans or self.current_type in ("hr",):
            el = {"type": self.current_type, "spans": spans}
            el.update(self.current_attrs)
            self.elements.append(el)
        self.current_type = None
        self.current_spans = []
        self.current_attrs = {}

    def _start_element(self, etype, **attrs):
        self._flush_element()
        self.current_type = etype
        self.current_attrs = attrs

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        # Headings
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._start_element(tag)

        # Paragraphs
        elif tag == "p":
            if self.in_blockquote:
                self._start_element("blockquote")
            else:
                self._start_element("p")

        # Lists
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
            # Mark this level as needing a counter restart on the next li
            self._restart_next_li.add(len(self.list_stack) - 1)
        elif tag == "li":
            list_type = self.list_stack[-1] if self.list_stack else "ul"
            level = len(self.list_stack) - 1
            restart = level in self._restart_next_li
            if restart:
                self._restart_next_li.discard(level)
            self._start_element("li", list_type=list_type, level=level, restart=restart)

        # Inline formatting
        elif tag in ("strong", "b"):
            self.format_stack.append("bold")
        elif tag in ("em", "i"):
            self.format_stack.append("italic")
        elif tag == "del":
            self.format_stack.append("strikethrough")
        elif tag == "mark":
            self.format_stack.append("mark")
        elif tag == "code" and not self.in_pre:
            self.format_stack.append("code")

        # Links
        elif tag == "a":
            self.current_href = attrs_dict.get("href", "")

        # Tables
        elif tag == "table":
            self._flush_element()
            self.in_table = True
            self.table_rows = []
        elif tag in ("tr", "thead", "tbody"):
            if tag == "tr":
                self.current_row = []
        elif tag in ("th", "td"):
            self.in_cell = True
            self.cell_tag = tag
            self.current_cell_spans = []

        # Blockquote
        elif tag == "blockquote":
            self.in_blockquote = True

        # Code blocks
        elif tag == "pre":
            self._flush_element()
            self.in_pre = True
            self.pre_text = []
        # code inside pre is just a marker, already handled by in_pre

        # Horizontal rule
        elif tag == "hr":
            self._flush_element()
            self.elements.append({"type": "hr", "spans": []})

        # Definition lists
        elif tag == "dl":
            self.in_dl = True
        elif tag == "dt":
            self._start_element("dt")
        elif tag == "dd":
            self._start_element("dd")

        # Line breaks
        elif tag == "br":
            if self.current_type is not None:
                self.current_spans.append({"text": "\n", **self._current_format_state()})
            elif self.in_cell:
                self.current_cell_spans.append({"text": "\n", **self._current_format_state()})

        # Images - skip with warning
        elif tag == "img":
            src = attrs_dict.get("src", "")
            alt = attrs_dict.get("alt", "")
            logger.warning(f"Image skipped: {src} (alt: {alt})")

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "dt", "dd"):
            self._flush_element()

        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()

        elif tag in ("strong", "b"):
            if "bold" in self.format_stack:
                self.format_stack.remove("bold")
        elif tag in ("em", "i"):
            if "italic" in self.format_stack:
                self.format_stack.remove("italic")
        elif tag == "del":
            if "strikethrough" in self.format_stack:
                self.format_stack.remove("strikethrough")
        elif tag == "mark":
            if "mark" in self.format_stack:
                self.format_stack.remove("mark")
        elif tag == "code" and not self.in_pre:
            if "code" in self.format_stack:
                self.format_stack.remove("code")

        elif tag == "a":
            self.current_href = None

        elif tag in ("th", "td"):
            self.in_cell = False
            cell_text = "".join(s["text"] for s in self.current_cell_spans)
            self.current_row.append({
                "text": cell_text,
                "spans": self.current_cell_spans,
                "header": self.cell_tag == "th",
            })
            self.current_cell_spans = []

        elif tag == "tr":
            if self.current_row:
                self.table_rows.append(self.current_row)
                self.current_row = []

        elif tag == "table":
            self.in_table = False
            if self.table_rows:
                self.elements.append({"type": "table", "rows": self.table_rows})
            self.table_rows = []

        elif tag == "blockquote":
            self.in_blockquote = False

        elif tag == "pre":
            self.in_pre = False
            code_text = "".join(self.pre_text)
            # Strip single trailing newline from code blocks
            if code_text.endswith("\n"):
                code_text = code_text[:-1]
            self.elements.append({"type": "code_block", "text": code_text, "spans": []})
            self.pre_text = []

        elif tag == "dl":
            self.in_dl = False

    def handle_data(self, data):
        # Code block - collect raw text
        if self.in_pre:
            self.pre_text.append(data)
            return

        # Table cell
        if self.in_cell:
            self.current_cell_spans.append({"text": data, **self._current_format_state()})
            return

        # Normal element span
        if self.current_type is not None and data:
            span = {"text": data, **self._current_format_state()}

            # Attach link href if inside <a>
            if self.current_href is not None:
                span["href"] = self.current_href

            self.current_spans.append(span)

    def get_elements(self):
        self._flush_element()
        return self.elements


# ---------------------------------------------------------------------------
# DOCX rendering functions
# ---------------------------------------------------------------------------

def lint_markdown(md_text):
    """Normalize markdown formatting using mdformat.
    Fixes blank lines before lists, consistent indentation, etc."""
    return mdformat.text(md_text)


def preprocess_markdown(md_text):
    """Pre-process markdown text before HTML conversion.
    Lints for consistent formatting, then handles Obsidian-specific syntax."""
    md_text = lint_markdown(md_text)
    # Convert ==highlight== to <mark> tags (after linting so mdformat doesn't touch them)
    md_text = re.sub(r'==(.*?)==', r'<mark>\1</mark>', md_text)
    return md_text


def parse_markdown_html(md_text):
    """Convert markdown text to HTML and parse into structured elements.
    Returns (elements, metadata_dict)."""
    md_text = preprocess_markdown(md_text)

    md_converter = markdown.Markdown(extensions=[
        'tables', 'fenced_code', 'footnotes',
        'meta', 'mdx_truly_sane_lists', 'smarty',
        'def_list', 'abbr',
    ])
    html = md_converter.convert(md_text)

    # Extract metadata from the meta extension
    metadata = {}
    if hasattr(md_converter, 'Meta'):
        for key, values in md_converter.Meta.items():
            metadata[key] = values[0] if len(values) == 1 else values

    parser = MarkdownHTMLParser()
    parser.feed(html)
    return parser.get_elements(), metadata


def add_formatted_spans(paragraph, spans, config):
    """Add a list of formatted spans to a paragraph as individual runs."""
    code_config = config.get("code", {})
    highlight_config = config.get("highlight", {})
    link_config = config.get("links", {})
    show_url = link_config.get("show_url", False)

    for span in spans:
        text = span["text"]
        href = span.get("href")

        # If link with show_url, append URL after text
        if href and show_url:
            text = f"{text} ({href})"

        run = paragraph.add_run(text)

        if span.get("bold"):
            run.bold = True
        if span.get("italic"):
            run.italic = True
        if span.get("strikethrough"):
            run.font.strike = True
        if span.get("code"):
            run.font.name = code_config.get("font_family", "Courier New")
            run.font.size = Pt(code_config.get("font_size", 10))
        if span.get("highlight"):
            run.bold = highlight_config.get("bold", True)
            # Yellow highlight via XML
            rPr = run._element.get_or_add_rPr()
            highlight_el = OxmlElement('w:highlight')
            highlight_el.set(qn('w:val'), 'yellow')
            rPr.append(highlight_el)


def apply_heading_style(paragraph, level, config):
    """Apply heading style to paragraph."""
    heading_key = f"h{level}"
    heading_config = config.get("headings", {}).get(heading_key, {})
    paragraph.style = f"Heading {level}"
    for run in paragraph.runs:
        if heading_config.get("size"):
            run.font.size = Pt(heading_config["size"])
        if heading_config.get("bold"):
            run.bold = True


def apply_body_style(paragraph, config):
    """Apply body text style to paragraph."""
    body_config = config.get("body", {})
    font_size = body_config.get("font_size", 11)
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
    if "paragraph_spacing_after" in body_config:
        paragraph.paragraph_format.space_after = Pt(body_config["paragraph_spacing_after"])


def _set_list_ilvl(paragraph, level):
    """Set the ilvl (indent level) on a list paragraph's numPr."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)
    ilvl_el = numPr.find(qn('w:ilvl'))
    if ilvl_el is None:
        ilvl_el = OxmlElement('w:ilvl')
        numPr.insert(0, ilvl_el)
    ilvl_el.set(qn('w:val'), str(level))


def _restart_list_numbering(doc, paragraph, level):
    """Create a new w:num entry that restarts at 1 for the given ilvl,
    and apply it to the paragraph. This ensures each new list block
    starts from 1 regardless of previous lists in the document."""
    try:
        np_part = doc.part.numbering_part
    except Exception:
        return

    numbering = np_part._element

    # Find the abstractNumId referenced by the paragraph's current numId.
    # First try direct numPr on paragraph, then fall back to style.
    current_num_id = None
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            numId_el = numPr.find(qn('w:numId'))
            if numId_el is not None:
                current_num_id = numId_el.get(qn('w:val'))

    if current_num_id is None:
        style = paragraph.style
        if style and style.element is not None:
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                s_numPr = s_pPr.find(qn('w:numPr'))
                if s_numPr is not None:
                    numId_el = s_numPr.find(qn('w:numId'))
                    if numId_el is not None:
                        current_num_id = numId_el.get(qn('w:val'))

    if current_num_id is None:
        return

    # Resolve to abstractNumId
    abstract_val = None
    for num_el in numbering.findall(qn('w:num')):
        if num_el.get(qn('w:numId')) == current_num_id:
            abstract_ref = num_el.find(qn('w:abstractNumId'))
            if abstract_ref is not None:
                abstract_val = abstract_ref.get(qn('w:val'))
            break

    if abstract_val is None:
        return

    # Assign a new numId value
    all_ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    new_num_id = max(all_ids) + 1 if all_ids else 1

    # Build <w:num w:numId="N"><w:abstractNumId w:val="X"/>
    #   <w:lvlOverride w:ilvl="L"><w:startOverride w:val="1"/></w:lvlOverride>
    # </w:num>
    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), str(new_num_id))

    abstract_ref_el = OxmlElement('w:abstractNumId')
    abstract_ref_el.set(qn('w:val'), abstract_val)
    new_num.append(abstract_ref_el)

    lvl_override = OxmlElement('w:lvlOverride')
    lvl_override.set(qn('w:ilvl'), str(level))
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    lvl_override.append(start_override)
    new_num.append(lvl_override)

    numbering.append(new_num)

    # Point the paragraph at this new numId
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)

    ilvl_el = numPr.find(qn('w:ilvl'))
    if ilvl_el is None:
        ilvl_el = OxmlElement('w:ilvl')
        numPr.insert(0, ilvl_el)
    ilvl_el.set(qn('w:val'), str(level))

    numId_el = numPr.find(qn('w:numId'))
    if numId_el is None:
        numId_el = OxmlElement('w:numId')
        numPr.append(numId_el)
    numId_el.set(qn('w:val'), str(new_num_id))


def add_code_block(doc, text, config):
    """Add a code block as a monospace paragraph with background shading."""
    code_config = config.get("code", {})
    font_family = code_config.get("font_family", "Courier New")
    font_size = code_config.get("font_size", 10)
    bg_color = code_config.get("background_color", "D9D9D9")

    for line in text.split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(line if line else " ")
        run.font.name = font_family
        run.font.size = Pt(font_size)

        # Background shading
        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        pPr.append(shd)


def add_table(doc, rows, config):
    """Add a table to the document."""
    if not rows:
        return
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_data in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                if cell_data.get("header"):
                    run = paragraph.add_run(cell_data["text"])
                    run.bold = True
                else:
                    add_formatted_spans(paragraph, cell_data.get("spans", []), config)


def add_horizontal_rule(doc):
    """Add a horizontal rule as a thin line."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_docx_metadata(doc, metadata):
    """Set DOCX core properties from frontmatter metadata."""
    props = doc.core_properties
    field_map = {
        "title": "title",
        "author": "author",
        "subject": "subject",
        "date": "created",
        "keywords": "keywords",
        "tags": "keywords",
    }
    for md_key, prop_name in field_map.items():
        value = metadata.get(md_key)
        if value:
            if isinstance(value, list):
                value = ", ".join(str(v) for v in value)
            else:
                value = str(value)
            try:
                setattr(props, prop_name, value)
            except (AttributeError, TypeError):
                logger.warning(f"Could not set DOCX property '{prop_name}' from '{md_key}'")


def convert_md_to_docx(md_text, config):
    """Convert markdown text to a DOCX Document object."""
    doc = Document()
    elements, metadata = parse_markdown_html(md_text)

    if metadata:
        set_docx_metadata(doc, metadata)

    list_config = config.get("lists", {})
    indent_size = list_config.get("indent_size", 0.5)
    bq_config = config.get("blockquote", {})

    for element in elements:
        etype = element["type"]

        if etype.startswith("h") and len(etype) == 2 and etype[1].isdigit():
            level = int(etype[1])
            paragraph = doc.add_paragraph()
            add_formatted_spans(paragraph, element["spans"], config)
            apply_heading_style(paragraph, level, config)

        elif etype == "p":
            paragraph = doc.add_paragraph()
            add_formatted_spans(paragraph, element["spans"], config)
            apply_body_style(paragraph, config)

        elif etype == "li":
            list_type = element.get("list_type", "ul")
            level = element.get("level", 0)
            restart = element.get("restart", False)
            style = f"List {'Number' if list_type == 'ol' else 'Bullet'}"
            paragraph = doc.add_paragraph(style=style)
            add_formatted_spans(paragraph, element["spans"], config)
            apply_body_style(paragraph, config)
            if list_type == "ol":
                if restart:
                    # Creates a new w:num with startOverride=1 at this ilvl
                    _restart_list_numbering(doc, paragraph, level)
                elif level > 0:
                    # Non-restart nested item: just set correct ilvl
                    _set_list_ilvl(paragraph, level)
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(indent_size * level)

        elif etype == "code_block":
            add_code_block(doc, element["text"], config)

        elif etype == "table":
            add_table(doc, element["rows"], config)

        elif etype == "hr":
            add_horizontal_rule(doc)

        elif etype == "blockquote":
            paragraph = doc.add_paragraph()
            add_formatted_spans(paragraph, element["spans"], config)
            bq_indent = bq_config.get("indent", 0.5)
            paragraph.paragraph_format.left_indent = Inches(bq_indent)
            if bq_config.get("italic", True):
                for run in paragraph.runs:
                    run.italic = True
            bq_font_size = bq_config.get("font_size", 11)
            for run in paragraph.runs:
                run.font.size = Pt(bq_font_size)

        elif etype == "dt":
            paragraph = doc.add_paragraph()
            add_formatted_spans(paragraph, element["spans"], config)
            for run in paragraph.runs:
                run.bold = True
            apply_body_style(paragraph, config)

        elif etype == "dd":
            paragraph = doc.add_paragraph()
            add_formatted_spans(paragraph, element["spans"], config)
            paragraph.paragraph_format.left_indent = Inches(indent_size)
            apply_body_style(paragraph, config)

    return doc


def save_docx(doc, output_path):
    """Save DOCX document to file."""
    doc.save(output_path)


def load_config(config_path=None):
    """Load styling configuration from YAML file."""
    if config_path is None:
        config_path = "config.yaml"

    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Config file not found: {config_path}")

    with open(config_path, "r") as f:
        config = yaml.safe_load(f)
    return config


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to DOCX format",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python md2docx.py input.md output.docx
  python md2docx.py input.md output.docx --config custom-config.yaml
        """,
    )

    parser.add_argument("input", help="Input markdown file")
    parser.add_argument("output", help="Output DOCX file")
    parser.add_argument(
        "--config", default="config.yaml", help="Config file (default: config.yaml)"
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    try:
        config = load_config(args.config)

        with open(args.input, "r", encoding="utf-8") as f:
            md_text = f.read()

        doc = convert_md_to_docx(md_text, config)
        save_docx(doc, args.output)

        print(f"Converted {args.input} -> {args.output}")
        sys.exit(0)

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except yaml.YAMLError as e:
        print(f"Error parsing config file: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


def lint_cli():
    """Standalone CLI for linting markdown files via mdformat."""
    parser = argparse.ArgumentParser(
        description="Lint/normalize markdown files using mdformat",
    )
    parser.add_argument("input", help="Input markdown file")
    parser.add_argument("-o", "--output", help="Output file (default: overwrite input)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()

    linted = lint_markdown(md_text)
    out_path = args.output or args.input
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(linted)

    print(f"Linted {args.input} -> {out_path}")


if __name__ == "__main__":
    main()
